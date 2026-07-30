"""
Build combined multi-panel figures and a Word document for journal submission review.

Produces:
  results/figures/final/fig1_combined.png   — Fig 1A (top) + Fig 1B (bottom)
  results/figures/final/fig3_combined.png   — Fig 3A (top) + Fig 3B (bottom)
  results/figures/final/fig4_combined.png   — Fig 4A (left) + Fig 4B (right)

  results/manuscript_figures.docx           — Word document, one figure per page,
                                              caption below, journal-submission layout

Run from project root:
    conda run -n eskape-ml python src/visualization/build_figures_docx.py
"""

import os
from pathlib import Path
from io import BytesIO

from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FINAL = Path("results/figures/final")
OUT_DIR = Path("results")
DPI = 300

# ── Layout helpers ────────────────────────────────────────────────────────────

def stack_vertical(paths: list[str], target_w_px: int = 1800, gap_px: int = 40) -> Image.Image:
    """Scale all images to target_w_px, then stack top-to-bottom with a white gap."""
    imgs = [Image.open(p).convert("RGB") for p in paths]
    resized = []
    for img in imgs:
        w, h = img.size
        new_h = round(h * target_w_px / w)
        resized.append(img.resize((target_w_px, new_h), Image.LANCZOS))
    total_h = sum(i.height for i in resized) + gap_px * (len(resized) - 1)
    canvas = Image.new("RGB", (target_w_px, total_h), "white")
    y = 0
    for img in resized:
        canvas.paste(img, (0, y))
        y += img.height + gap_px
    return canvas


def stack_horizontal(paths: list[str], target_h_px: int = 1200, gap_px: int = 40) -> Image.Image:
    """Scale all images to target_h_px, then place left-to-right with a white gap."""
    imgs = [Image.open(p).convert("RGB") for p in paths]
    resized = []
    for img in imgs:
        w, h = img.size
        new_w = round(w * target_h_px / h)
        resized.append(img.resize((new_w, target_h_px), Image.LANCZOS))
    total_w = sum(i.width for i in resized) + gap_px * (len(resized) - 1)
    canvas = Image.new("RGB", (total_w, target_h_px), "white")
    x = 0
    for img in resized:
        canvas.paste(img, (x, 0))
        x += img.width + gap_px
    return canvas


def save_png(img: Image.Image, path: Path) -> None:
    img.save(path, format="PNG", dpi=(DPI, DPI))
    print(f"Saved: {path}")


# ── Build combined figures ────────────────────────────────────────────────────

# Fig 1: three panels in one row — A (confusion matrix), B (holdout recall),
# C (holdout BA vs CV). All panels are ~1300px tall; scale to common height.
fig1 = stack_horizontal([
    "results/figures/rf/fig1a_q1_confusion_matrix.png",
    "results/figures/rf/fig1b_recall.png",
    "results/figures/rf/fig1c_ba_comparison.png",
], target_h_px=1300)
save_png(fig1, FINAL / "fig1_combined.png")

# Fig 2: Q2 ARG burden — forest plot (top) + driver dotplot (bottom)
fig2 = stack_vertical([
    FINAL / "fig2a_q2_auroc_forest.png",
    FINAL / "fig2b_q2_driver_dotplot.png",
], target_w_px=2400)
save_png(fig2, FINAL / "fig2_combined.png")

# Fig 4: SHAP — beeswarm (left) + heatmap (right)
fig4 = stack_horizontal([
    FINAL / "fig4a_shap_global_beeswarm.png",
    FINAL / "fig4b_shap_heatmap.png",
], target_h_px=1800)
save_png(fig4, FINAL / "fig4_combined.png")

print()


# ── Word document ─────────────────────────────────────────────────────────────

# mBio-style: A4, 2.54 cm margins all round; 11 pt Times New Roman body
doc = Document()

# Page layout: A4, 2.54 cm margins
sec = doc.sections[0]
sec.page_height = Cm(29.7)
sec.page_width  = Cm(21.0)
for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, attr, Cm(2.54))

# Text width available: 21 - 2*2.54 = 15.92 cm ≈ 6.27 in
TEXT_W = Inches(6.27)

# Default paragraph style
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def add_figure(doc: Document, img_path: Path, caption_title: str, caption_body: str) -> None:
    """Add one figure (image + caption) on a new page."""
    # Centred image paragraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=TEXT_W)

    # Small vertical spacer
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(6)
    spacer.paragraph_format.space_after  = Pt(0)

    # Caption title (bold) + body (normal), both 10 pt
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap_para.paragraph_format.space_before = Pt(0)
    cap_para.paragraph_format.space_after  = Pt(12)

    title_run = cap_para.add_run(caption_title + " ")
    title_run.bold = True
    title_run.font.size = Pt(10)
    title_run.font.name = "Times New Roman"

    body_run = cap_para.add_run(caption_body)
    body_run.bold = False
    body_run.font.size = Pt(10)
    body_run.font.name = "Times New Roman"


# ── Title ─────────────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("Figures — ESKAPE Defence Systems ML Study")
title_run.bold = True
title_run.font.size = Pt(14)
title_run.font.name = "Times New Roman"
doc.add_paragraph()  # blank line

# ── Figure 1 ──────────────────────────────────────────────────────────────────
page_break(doc)
add_figure(
    doc,
    FINAL / "fig1_combined.png",
    "FIG 1. Species classification performance of the defence-system Random Forest classifier.",
    (
        "(A) Row-normalised confusion matrix from phylogenetically grouped five-fold "
        "cross-validation (n = 3,335 genomes; 309 phylogroups; 359 defence system binary "
        "features). Diagonal values are per-species recall. Balanced accuracy = 0.900 "
        "[95% CI: 0.871–0.923]; macro-F1 = 0.896 [0.843–0.949]. "
        "S. aureus recall = 1.000; E. faecium = 0.985; "
        "K. pneumoniae = 0.909; E. cloacae complex = 0.901; "
        "P. aeruginosa = 0.893; A. baumannii = 0.683 "
        "(23% of AB genomes misclassified as E. faecium; see Results §1 for mechanism). "
        "Class abbreviations: AB, A. baumannii; EC, E. cloacae complex; EF, E. faecium; "
        "KP, K. pneumoniae; PA, P. aeruginosa; SA, S. aureus. "
        "(B) Per-species recall on the C3 holdout set (180 NCBI complete genomes withheld "
        "entirely from training; 30 per ESKAPE species; novel-ST-prioritised; "
        "Supplementary Table S16). "
        "(C) Holdout balanced accuracy (0.944 [0.911–0.972]) plotted alongside the grouped "
        "cross-validation estimate (0.900 [0.871–0.923]), confirming generalisation beyond "
        "the training cohort. Horizontal bars show 95% CIs. Error bar on the CV estimate is "
        "the 95% CI from phylogroup cluster bootstrap (2,000 iterations). Error bar on the "
        "holdout estimate is the 95% CI from bootstrap resampling of the 180 holdout genomes."
    ),
)

# ── Figure 2 ──────────────────────────────────────────────────────────────────
page_break(doc)
add_figure(
    doc,
    FINAL / "fig2_combined.png",
    "FIG 2. Association between defence system profiles and ARG burden across ESKAPE species.",
    (
        "(A) Per-species AUROC from the Q2 ARG burden classifier (Random Forest, binary high "
        "vs low ARG tertile, phylogenetically grouped five-fold cross-validation). Error bars "
        "are 95% CIs from fold-level bootstrap (2,000 iterations). BH-adjusted significance "
        "(one-sample t-test against null BA = 0.5; q = 0.05) was confirmed "
        "in four species (EC, EF, KP, PA) and not confirmed in two (AB, SA; grey labels). "
        "AUROC was 0.824 [0.775\u20130.872] for EF (p_adj = 0.0082), "
        "0.791 [0.688\u20130.894] for KP (p_adj = 0.038), 0.791 [0.701\u20130.870] "
        "for PA (p_adj = 0.038), 0.781 [0.704\u20130.843] for EC "
        "(p_adj = 0.0012), 0.757 [0.659\u20130.854] for AB (ns), and "
        "0.724 [0.596\u20130.814] for SA (ns). The null (AUROC = 0.5) is shown as "
        "a dashed vertical line. "
        "(B) Within-phylogroup-robust ARG-burden drivers for the four Q2-significant species, "
        "ordered by within-phylogroup Spearman \u03c1 descending. Each point represents a "
        "feature that survived the within-phylogroup gate (p_within_pg < 0.05 with "
        "sign preserved). Positive \u03c1 denotes a facilitative association (higher defence "
        "system presence in high-ARG genomes within the same phylogroup). Negative \u03c1 "
        "denotes a restrictive association. Blue dots indicate named systems with mechanistically "
        "characterised functions. Orange dots indicate uncharacterised candidate systems. "
        "Panel subtitles show AUROC, the count of all robust drivers, and the number shown "
        "(capped at 8 per panel for readability). Full driver lists in Supplementary Tables "
        "S2 and S13."
    ),
)

# ── Figure 3 ──────────────────────────────────────────────────────────────────
page_break(doc)
add_figure(
    doc,
    FINAL / "fig3_q3b_block_comparison.png",
    "FIG 3. Species-discriminating power of four genomic feature blocks after near-exclusive marker removal.",
    (
        "Bar chart of Adjusted Rand Index (ARI) between K-means partitions (K = 6) "
        "and species identity labels, computed on 309 dereplicated phylogroup representatives. "
        "The five conditions were defence only (359 features, identical to the primary Q3 "
        "analysis), IS element family presence (19 families after filtering), HMRG class "
        "presence (9 classes after filtering), ARG gene presence (134 genes after filtering), "
        "and Defence combined with IS elements (378 features after filtering). Near-exclusive "
        "species markers (spec_score \u2265 0.70) were removed from IS, ARG, and "
        "HMRG blocks before clustering to prevent taxonomically-concentrated features from "
        "inflating ARI. The same threshold was already applied to defence features in the "
        "primary analysis. Ghost bars (hatched outlines) show pre-filter ARI for each block, "
        "quantifying the marker-driven inflation. Error bars are 95% CIs from 2,000-iteration "
        "bootstrap resampling of the 309 genomes. All conditions were significant by permutation "
        "test (p = 0.002, 500 permutations, K = 6). The dashed blue line "
        "marks the Defence-only baseline (ARI = 0.219). \u2020 denotes blocks for "
        "which near-exclusive marker removal was applied."
    ),
)

# ── Figure 4 ──────────────────────────────────────────────────────────────────
page_break(doc)
add_figure(
    doc,
    FINAL / "fig4_combined.png",
    "FIG 4. SHAP attribution of defence-system features to the Q1 species classifier.",
    (
        "(A) Global SHAP beeswarm for the top 20 features by mean absolute SHAP value "
        "(averaged over all 3,335 genomes and all six species classes). Each row is one "
        "feature and each dot is one genome. Dot position on the x-axis is the mean absolute "
        "SHAP value (mean |SHAP| over the six class outputs) for that genome and feature. "
        "Dot colour encodes feature value, with orange-red indicating presence (value = 1) "
        "and blue indicating absence (value = 0). Features are ordered by descending "
        "global mean |SHAP|. All features shown are binary (defence system present/absent per "
        "genome). The x-axis reflects importance magnitude. Signed per-class effects are shown "
        "in panel B. "
        "(B) Per-species conditional mean SHAP for the global top 20 features. Heatmap of mean "
        "signed SHAP value among genomes where the feature is present. Orange indicates that "
        "feature presence increases model probability toward that species. Blue indicates that "
        "feature presence decreases model probability toward that species. Rows are ordered by "
        "global rank (rank 1 at top). Column abbreviations as in Fig. 1. Row-rank numbers "
        "(#1 to #20) on the right axis. \u2020 marks uncharacterised candidate systems. "
        "Because the heatmap conditions on feature presence, the magnitudes differ from panel A "
        "(which averages over all genomes including those where the feature is absent)."
    ),
)

# ── Supplementary Figure S-A ──────────────────────────────────────────────────
page_break(doc)
add_figure(
    doc,
    FINAL / "sfig_sa_shap_within_species.png",
    "Supplementary Figure S-A. Within-species SHAP distributions for top-ranked defence system features.",
    (
        "Violin plots of SHAP values (for the species-own class) for the top six features by "
        "mean absolute SHAP within each ESKAPE species class (Q1 RF). Plots are split by "
        "feature value, with blue indicating genomes in which the feature is present and grey "
        "indicating genomes in which the feature is absent. The dashed vertical line marks 0 "
        "(no contribution to class probability). Inner quartile lines show the 25th, 50th, and "
        "75th percentiles. Features are excluded if either presence or absence group contains "
        "fewer than 5 genomes in that species. Panel subtitle indicates the number of top-six "
        "features excluded by this criterion (shown in grey in the title line). These plots "
        "show signed effects within the species of interest. Global signed effects across all "
        "species classes are shown in Fig. 4B."
    ),
)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = OUT_DIR / "manuscript_figures.docx"
doc.save(str(out_path))
print(f"Saved: {out_path}")
